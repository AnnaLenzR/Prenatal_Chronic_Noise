import csv
import re
import shutil
import sys
from collections import Counter
from datetime import date
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/annalenz/Desktop/ChronicNoise_Rodents/Prenatal_Chronic_Noise")
REFERENCE = Path("/Users/annalenz/Desktop/Bird_map/outlines/Bird_noise_map_screening_counts.docx")
OUTPUT_DIR = ROOT / "outputs" / "prisma"
FIGURE_DIR = ROOT / "Figures"
DOCX_OUT = OUTPUT_DIR / "PCN_PRISMA_screening_counts.docx"
PNG_OUT = FIGURE_DIR / "PCN_PRISMA_flowchart.png"
FULL_TEXT_CSV = ROOT / "PCN_screening" / "full_text_2026-08-28_23-47-26" / "articles.csv"
DATA_CSV = ROOT / "Data" / "PCN_data_ext_checking_v.3.csv"

FONT_REGULAR = "/System/Library/Fonts/Supplemental/Arial.ttf"
FONT_BOLD = "/System/Library/Fonts/Supplemental/Arial Bold.ttf"

NAVY = "17365D"
BLUE = "2E74B5"
TEAL = "13698B"
LIGHT_BLUE = "D9E8F6"
PALE_BLUE = "EEF3F7"
PEACH = "F7D1B0"
PALE_PEACH = "FBE2D3"
GRID = "AFC3D4"
GRAY = "5F6F7D"
WHITE = "FFFFFF"
BLACK = "172B3A"


SOURCE_COUNTS = [
    ("Scopus", 1404),
    ("Web of Science", 1842),
    ("PubMed", 1056),
    ("PsycINFO", 408),
    ("OpenAlex", 681),
    ("Google Scholar - supplementary search", 489),
    ("BASE", 9),
]

MISSING_REASON_MAP = {
    "386282729": "Wrong exposure",
    "386283068": "Wrong population",
    "386283297": "Wrong exposure",
    "386283305": "Foreign language",
    "386283326": "Wrong exposure",
    "386283372": "Wrong exposure",
}

REASON_LABELS = {
    "wrong exposure": "Wrong exposure",
    "wrong outcome": "Wrong outcome",
    "wrong publication type": "Wrong publication type",
    "wrong population": "Wrong population",
    "foreign language": "Foreign language",
    "full text not available": "Full text not available",
}


def parse_decision(note):
    match = re.search(r'RAYYAN-INCLUSION: \{"Anna"=>"([^"]+)"\}', note or "")
    return match.group(1) if match else "Missing"


def parse_reasons(note):
    match = re.search(r"RAYYAN-EXCLUSION-REASONS: (.*?)(?: \| [A-Z][A-Z-]+:|$)", note or "")
    if not match:
        return []
    return [part.strip().lower() for part in match.group(1).split(",") if part.strip()]


def clean_title(raw):
    return re.sub(r"\s+SO\s+-\s+", ". ", (raw or "").strip()).replace("…", "...")


def reference_text(row):
    title = clean_title(row.get("title", ""))
    year = (row.get("year") or "").strip()
    doi = (row.get("doi") or "").strip()
    parts = [title]
    if year and f"({year})" not in title:
        parts.append(f"({year})")
    if doi and doi not in title:
        doi = doi.replace("https://doi.org/", "")
        parts.append(f"DOI: {doi}")
    return ". ".join(p.strip(". ") for p in parts if p).strip() + "."


def load_full_text_records():
    with FULL_TEXT_CSV.open(encoding="utf-8-sig", newline="") as handle:
        rows = list(csv.DictReader(handle))

    excluded = []
    not_retrieved = []
    included = []
    for row in rows:
        article_id = row["key"].removeprefix("rayyan-")
        decision = parse_decision(row.get("notes", ""))
        reasons = parse_reasons(row.get("notes", ""))
        if decision == "Included":
            included.append(row)
            continue
        if decision != "Excluded":
            raise AssertionError(f"Unexpected decision for {article_id}: {decision}")

        if "full text not available" in reasons:
            reason = "Full text not available"
        elif "foreign language" in reasons:
            reason = "Foreign language"
        elif reasons:
            reason = REASON_LABELS[reasons[0]]
        else:
            reason = MISSING_REASON_MAP[article_id]

        record = {
            "rayyan_id": article_id,
            "reference": reference_text(row),
            "reason": reason,
        }
        if reason == "Full text not available":
            not_retrieved.append(record)
        else:
            excluded.append(record)

    reason_counts = Counter(item["reason"] for item in excluded)
    expected = {
        "Wrong exposure": 60,
        "Wrong outcome": 4,
        "Wrong publication type": 4,
        "Wrong population": 4,
        "Foreign language": 2,
    }
    assert len(rows) == 87, len(rows)
    assert len(included) == 12, len(included)
    assert len(excluded) == 74, len(excluded)
    assert len(not_retrieved) == 1, len(not_retrieved)
    assert dict(reason_counts) == expected, (reason_counts, expected)
    return rows, included, excluded, not_retrieved, expected


def load_analysis_studies():
    with DATA_CSV.open(encoding="utf-8-sig", newline="") as handle:
        rows = list(csv.DictReader(handle))
    studies = {}
    for row in rows:
        studies.setdefault(row["study_id"], {"title": " ".join(row["title"].split()), "doi": row["doi"].strip()})
    assert len(studies) == 16, len(studies)
    return studies


def wrap_lines(draw, text, font, max_width):
    lines = []
    for paragraph in str(text).split("\n"):
        words = paragraph.split()
        if not words:
            lines.append("")
            continue
        current = words[0]
        for word in words[1:]:
            trial = current + " " + word
            if draw.textlength(trial, font=font) <= max_width:
                current = trial
            else:
                lines.append(current)
                current = word
        lines.append(current)
    return lines


def draw_centered_text(draw, box, text, font, fill, spacing=8, max_width_pad=34):
    x1, y1, x2, y2 = box
    lines = wrap_lines(draw, text, font, (x2 - x1) - 2 * max_width_pad)
    bbox = draw.multiline_textbbox((0, 0), "\n".join(lines), font=font, spacing=spacing, align="center")
    height = bbox[3] - bbox[1]
    y = y1 + ((y2 - y1) - height) / 2 - bbox[1]
    draw.multiline_text(((x1 + x2) / 2, y), "\n".join(lines), font=font, fill=fill, spacing=spacing, align="center", anchor="ma")


def rounded_box(draw, box, fill, outline, width=7, radius=18):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def arrow(draw, points, color, width=8, head=25):
    draw.line(points, fill=color, width=width, joint="curve")
    x1, y1 = points[-2]
    x2, y2 = points[-1]
    if abs(x2 - x1) >= abs(y2 - y1):
        if x2 > x1:
            poly = [(x2, y2), (x2 - head, y2 - head * 0.65), (x2 - head, y2 + head * 0.65)]
        else:
            poly = [(x2, y2), (x2 + head, y2 - head * 0.65), (x2 + head, y2 + head * 0.65)]
    else:
        if y2 > y1:
            poly = [(x2, y2), (x2 - head * 0.65, y2 - head), (x2 + head * 0.65, y2 - head)]
        else:
            poly = [(x2, y2), (x2 - head * 0.65, y2 + head), (x2 + head * 0.65, y2 + head)]
    draw.polygon(poly, fill=color)


def create_flowchart(path):
    width, height = 4500, 3300
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    f_band = ImageFont.truetype(FONT_BOLD, 52)
    f_source = ImageFont.truetype(FONT_BOLD, 40)
    f_main = ImageFont.truetype(FONT_BOLD, 48)
    f_right = ImageFont.truetype(FONT_REGULAR, 38)
    f_right_bold = ImageFont.truetype(FONT_BOLD, 38)
    f_reason = ImageFont.truetype(FONT_REGULAR, 34)
    f_rail = ImageFont.truetype(FONT_BOLD, 39)

    blue = "#" + TEAL
    dark = "#" + NAVY
    pale = "#" + PALE_BLUE
    peach = "#" + PEACH
    lightblue = "#" + LIGHT_BLUE

    draw.rounded_rectangle((270, 65, 4380, 175), radius=38, fill=lightblue)
    draw_centered_text(draw, (270, 65, 4380, 175), "Records identified from databases and search sources", f_band, blue)

    rails = [
        ((45, 210, 190, 1100), "Identification"),
        ((45, 1140, 190, 2390), "Screening"),
        ((45, 2430, 190, 3210), "Included"),
    ]
    for box, label in rails:
        draw.rounded_rectangle(box, radius=65, fill=peach)
        temp = Image.new("RGBA", (box[3] - box[1], box[2] - box[0]), (255, 255, 255, 0))
        td = ImageDraw.Draw(temp)
        td.text((temp.width / 2, temp.height / 2), label, font=f_rail, fill="#9B4A18", anchor="mm")
        temp = temp.rotate(90, expand=True)
        image.paste(temp, (int((box[0] + box[2] - temp.width) / 2), int((box[1] + box[3] - temp.height) / 2)), temp)

    top_source_boxes = [
        ((280, 230, 1220, 435), "Scopus\n(n = 1,404)"),
        ((1290, 230, 2230, 435), "Web of Science\n(n = 1,842)"),
        ((2300, 230, 3240, 435), "PubMed\n(n = 1,056)"),
        ((3310, 230, 4250, 435), "PsycINFO\n(n = 408)"),
    ]
    lower_source_boxes = [
        ((350, 520, 1450, 725), "OpenAlex\n(n = 681)"),
        ((1700, 520, 2800, 725), "Google Scholar\nsupplementary search\n(n = 489)"),
        ((3050, 520, 4150, 725), "BASE\n(n = 9)"),
    ]
    source_boxes = top_source_boxes + lower_source_boxes
    for box, text in source_boxes:
        rounded_box(draw, box, "white", blue, width=7, radius=4)
        draw_centered_text(draw, box, text, f_source, dark, spacing=5)

    # Source aggregation lines. The two rows use separate trunks so no
    # connector runs through a source box.
    upper_trunk_y = 475
    upper_centres = []
    for box, _ in top_source_boxes:
        cx = (box[0] + box[2]) / 2
        upper_centres.append(cx)
        draw.line((cx, box[3], cx, upper_trunk_y), fill=blue, width=6)
    draw.line((min(upper_centres), upper_trunk_y, max(upper_centres), upper_trunk_y), fill=blue, width=6)

    merge_x = 1550
    lower_trunk_y = 790
    draw.line((merge_x, upper_trunk_y, merge_x, lower_trunk_y), fill=blue, width=6)
    lower_centres = []
    for box, _ in lower_source_boxes:
        cx = (box[0] + box[2]) / 2
        lower_centres.append(cx)
        draw.line((cx, box[3], cx, lower_trunk_y), fill=blue, width=6)
    draw.line((min(lower_centres), lower_trunk_y, max(lower_centres), lower_trunk_y), fill=blue, width=6)

    records = (620, 870, 2080, 1050)
    duplicates = (2460, 850, 4260, 1090)
    rounded_box(draw, records, "white", blue)
    draw_centered_text(draw, records, "Records identified\n(n = 5,889)", f_main, dark)
    arrow(draw, [(merge_x, lower_trunk_y), (merge_x, 825), (1350, 825), (1350, records[1])], blue)
    rounded_box(draw, duplicates, pale, blue)
    draw.text((2510, 895), "Duplicates removed before screening (n = 3,576)", font=f_right_bold, fill=dark)
    draw.text((2510, 960), "Includes duplicate BASE records", font=f_right, fill=dark)
    arrow(draw, [(records[2], 960), (duplicates[0], 960)], blue)

    unique = (620, 1200, 2080, 1380)
    rounded_box(draw, unique, "white", blue)
    draw_centered_text(draw, unique, "Unique records screened\n(n = 2,313)", f_main, dark)
    arrow(draw, [(1350, records[3]), (1350, unique[1])], blue)

    formal = (620, 1540, 2080, 1735)
    rounded_box(draw, formal, "white", blue)
    draw_centered_text(draw, formal, "Included at title-and-\nabstract stage\n(n = 87)", f_main, dark, spacing=4)
    arrow(draw, [(1350, unique[3]), (1350, formal[1])], blue)

    ta_excl = (2460, 1515, 4260, 1760)
    rounded_box(draw, ta_excl, pale, blue)
    draw_centered_text(draw, ta_excl, "Records excluded after title-and-abstract screening\n(n = 2,226)", f_right, dark)
    arrow(draw, [(formal[2], 1638), (ta_excl[0], 1638)], blue)

    rayyan_included = (620, 2070, 2080, 2260)
    rounded_box(draw, rayyan_included, "white", blue)
    draw_centered_text(draw, rayyan_included, "Included at full-text stage\n(n = 12)", f_main, dark, spacing=4)
    arrow(draw, [(1350, formal[3]), (1350, rayyan_included[1])], blue)

    excluded = (2460, 1940, 4260, 2390)
    rounded_box(draw, excluded, pale, blue)
    draw.text((2510, 1980), "Full-text reports excluded, with reasons (n = 75)", font=f_right_bold, fill=dark)
    reasons = [
        "Wrong exposure (n = 60)",
        "Wrong outcome (n = 4)",
        "Wrong publication type (n = 4)",
        "Wrong population (n = 4)",
        "Foreign language (n = 2)",
        "Full text not available (n = 1)",
    ]
    for idx, line in enumerate(reasons):
        draw.text((2510, 2050 + idx * 52), line, font=f_reason, fill=dark)
    arrow(draw, [(rayyan_included[2], 2165), (excluded[0], 2165)], blue)

    additional = (2460, 2500, 4260, 2710)
    rounded_box(draw, additional, "#" + PALE_PEACH, blue)
    draw_centered_text(draw, additional, "Other sources (benchmark set)\n(n = 4)", f_main, dark, spacing=4)

    final = (1150, 2920, 2750, 3150)
    rounded_box(draw, final, peach, blue)
    draw_centered_text(draw, final, "Studies included in\nthe meta-analysis\n(n = 16)", f_main, dark, spacing=4)
    arrow(draw, [(1350, rayyan_included[3]), (1350, 2800), (1750, 2800), (1750, final[1])], blue)
    arrow(draw, [(3360, additional[3]), (3360, 2800), (2350, 2800), (2350, final[1])], blue)

    path.parent.mkdir(parents=True, exist_ok=True)
    image.save(path, quality=95, dpi=(300, 300))


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, *, bold=False, color=BLACK, size=10, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table, color=GRID, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(GRAY)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def set_header_footer(section, left_text, right_text):
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.text = ""
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(6.45), 2)
    left = paragraph.add_run(left_text.upper())
    left.bold = True
    left.font.name = "Calibri"
    left.font.size = Pt(8)
    left.font.color.rgb = RGBColor.from_string(GRAY)
    paragraph.add_run("\t")
    right = paragraph.add_run(right_text.upper())
    right.font.name = "Calibri"
    right.font.size = Pt(8)
    right.font.color.rgb = RGBColor.from_string(GRAY)
    footer_p = section.footer.paragraphs[0]
    footer_p.text = ""
    add_page_number(footer_p)


def remove_document_body(doc):
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_callout(doc, label, text, table_geometry):
    table = doc.add_table(rows=1, cols=1)
    repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run(label + ": ")
    r1.bold = True
    r1.font.color.rgb = RGBColor.from_string(NAVY)
    r2 = p.add_run(text)
    for run in (r1, r2):
        run.font.name = "Calibri"
        run.font.size = Pt(10.5)
    set_cell_shading(cell, "F2F5F8")
    set_table_borders(table, "CAD8E4", "5")
    table_geometry(table, [9360], table_width_dxa=9360)
    return table


def add_data_table(doc, headers, rows, widths_dxa, table_geometry, *, appendix=False, total_last=False):
    table = doc.add_table(rows=1, cols=len(headers))
    for col, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[col], header, bold=True, color=NAVY, size=9.5 if appendix else 10)
        set_cell_shading(table.rows[0].cells[col], "DCE6F1")
    repeat_table_header(table.rows[0])
    for row_index, row_data in enumerate(rows):
        row = table.add_row()
        for col, value in enumerate(row_data):
            align = WD_ALIGN_PARAGRAPH.RIGHT if (not appendix and col == 1 and isinstance(value, int)) else WD_ALIGN_PARAGRAPH.LEFT
            if appendix and col == 0:
                align = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(row.cells[col], f"{value:,}" if isinstance(value, int) else value, size=8.5 if appendix else 10, align=align)
        prevent_row_split(row)
        if total_last and row_index == len(rows) - 1:
            for cell in row.cells:
                set_cell_shading(cell, "EEF2F6")
                for run in cell.paragraphs[0].runs:
                    run.bold = True
    set_table_borders(table)
    table_geometry(table, widths_dxa, table_width_dxa=sum(widths_dxa))
    return table


def create_docx(path, flowchart_path, included, excluded, not_retrieved, reason_counts, studies):
    sys.path.insert(0, "/Users/annalenz/.codex/plugins/cache/openai-primary-runtime/documents/26.826.12353/skills/documents/scripts")
    from table_geometry import apply_table_geometry

    path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(REFERENCE, path)
    doc = Document(path)
    remove_document_body(doc)

    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    set_header_footer(section, "Prenatal chronic noise meta-analysis", "Working record")

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)

    title = doc.add_paragraph(style="Title")
    title.add_run("Evidence-selection counts")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("Prenatal chronic noise exposure in rodents meta-analysis")
    meta = doc.add_paragraph()
    meta_run = meta.add_run("Working document | 28 August 2026")
    meta_run.italic = True
    meta_run.font.color.rgb = RGBColor.from_string(GRAY)

    add_callout(
        doc,
        "Scope",
        "Identification and screening counts use records/reports; the final inclusion count uses studies represented in the meta-analysis dataset.",
        apply_table_geometry,
    )

    add_heading(doc, "Records identified by database or search source")
    source_rows = [(name, count) for name, count in SOURCE_COUNTS] + [("Total records identified", 5889)]
    add_data_table(doc, ["Database or search source", "Records"], source_rows, [7560, 1800], apply_table_geometry, total_last=True)
    check = doc.add_paragraph("Check: 1,404 + 1,842 + 1,056 + 408 + 681 + 489 + 9 = 5,889.")
    check.runs[0].italic = True
    check.runs[0].font.color.rgb = RGBColor.from_string(GRAY)

    add_heading(doc, "Deduplication, screening, and final inclusion")
    stage_rows = [
        ("Records identified", 5889, "Sum across all database and search-source results"),
        ("Duplicates removed before screening", 3576, "Combined deduplication, including repeated BASE records"),
        ("Unique records screened", 2313, "5,889 - 3,576"),
        ("Title-and-abstract exclusions", 2226, "Reasons are not reported at this stage"),
        ("Included at title-and-abstract stage", 87, "2,313 - 2,226"),
        ("Full-text reports excluded", 75, "Includes one report for which full text was unavailable"),
        ("Included at full-text stage", 12, "87 - 75"),
        ("Other sources (benchmark set)", 4, "Added outside the formal search strategy"),
        ("Studies included in the meta-analysis", 16, "12 + 4"),
    ]
    add_data_table(doc, ["Selection stage", "Records", "Reconciliation note"], stage_rows, [4740, 1260, 3360], apply_table_geometry)

    add_heading(doc, "Full-text exclusion reasons")
    reason_order = ["Wrong exposure", "Wrong outcome", "Wrong publication type", "Wrong population", "Foreign language"]
    reason_rows = [(reason, reason_counts[reason]) for reason in reason_order]
    reason_rows += [("Full text not available", 1), ("Total full-text exclusions", 75)]
    add_data_table(doc, ["Primary exclusion reason", "Reports"], reason_rows, [7560, 1800], apply_table_geometry, total_last=True)
    note = doc.add_paragraph()
    nr = note.add_run("Foreign language received priority when a record had multiple active Rayyan labels. The report whose full text was unavailable is included as an exclusion reason, as requested.")
    nr.italic = True
    nr.font.color.rgb = RGBColor.from_string(GRAY)

    add_heading(doc, "Arithmetic reconciliation")
    arithmetic = [
        "5,889 - 3,576 = 2,313 unique records screened",
        "2,313 - 2,226 = 87 included at the title-and-abstract stage",
        "87 - 75 = 12 included at the full-text stage",
        "12 + 4 = 16 studies included in the meta-analysis",
    ]
    for line in arithmetic:
        p = doc.add_paragraph(style="List Number")
        p.add_run(line)

    add_heading(doc, "Source files")
    source_files = doc.add_paragraph()
    source_files.paragraph_format.space_after = Pt(0)
    source_files.paragraph_format.line_spacing = 1
    source_files.add_run(
        "Title/abstract: PCN_screening/title_abstract_2026-08-28_23-45-55/articles.csv\n"
        "Full text: PCN_screening/full_text_2026-08-28_23-47-26/articles.csv\n"
        "Additional reports: PCN_screening/Benchmark/\n"
        "Included-study data: Data/PCN_data_ext_checking_v.3.csv"
    )
    for run in source_files.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(GRAY)

    landscape = doc.add_section(WD_SECTION.NEW_PAGE)
    landscape.orientation = WD_ORIENT.LANDSCAPE
    landscape.page_width = Inches(11)
    landscape.page_height = Inches(8.5)
    landscape.left_margin = Inches(0.75)
    landscape.right_margin = Inches(0.75)
    landscape.top_margin = Inches(0.75)
    landscape.bottom_margin = Inches(0.75)
    set_header_footer(landscape, "Prenatal chronic noise meta-analysis", "Full-text decision appendix")

    add_heading(doc, "Appendix A. Reports excluded at the full-text stage (n = 75)")
    intro = doc.add_paragraph("Each report has one primary exclusion reason. Foreign language received priority when a report had multiple active labels. The unavailable full text is included in this appendix as an exclusion reason.")
    intro.runs[0].font.color.rgb = RGBColor.from_string(GRAY)
    combined_exclusions = excluded + [dict(item, reason="Full text not available") for item in not_retrieved]
    appendix_rows = [(item["rayyan_id"], item["reference"], item["reason"]) for item in sorted(combined_exclusions, key=lambda x: (x["reason"], int(x["rayyan_id"])))]
    add_data_table(doc, ["Rayyan ID", "Reference", "Primary reason"], appendix_rows, [1220, 10060, 2400], apply_table_geometry, appendix=True)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    add_heading(doc, "Appendix B. Reports included from other sources (benchmark set) (n = 4)")
    p = doc.add_paragraph("These four reports came from the benchmark set and were added to the final included-study set. They were not counted among the 87 Rayyan full-text records.")
    p.runs[0].font.color.rgb = RGBColor.from_string(GRAY)
    additional_ids = [
        ("abramova_2021_front", "Benchmark set"),
        ("arjunan_2023_stress", "Benchmark set"),
        ("oliveira_2015_jbehavbraisci", "Benchmark set"),
        ("abramova_2023_biopsy", "Benchmark set"),
    ]
    additional_rows = []
    for study_id, source in additional_ids:
        study = studies[study_id]
        ref = study["title"]
        if study["doi"]:
            ref += f". DOI: {study['doi'].replace('https://doi.org/', '')}"
        additional_rows.append((study_id, ref + ".", source))
    add_data_table(doc, ["Study ID", "Reference", "Identification route"], additional_rows, [2350, 8720, 2610], apply_table_geometry, appendix=True)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    add_heading(doc, "Identification and selection of prenatal chronic-noise studies")
    picture = doc.add_picture(str(flowchart_path), width=Inches(8.2))
    picture._inline.docPr.set(
        "descr",
        "PRISMA-style flow diagram showing 5,889 identified records, 3,576 duplicates, 2,313 screened records, 75 full-text exclusions, and 16 included studies.",
    )
    picture._inline.docPr.set("title", "PRISMA flow diagram")
    picture_paragraph = doc.paragraphs[-1]
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_paragraph.paragraph_format.space_after = Pt(3)
    caption = doc.add_paragraph("Figure 1. PRISMA-style flow diagram for identification, screening, full-text assessment, and inclusion in the meta-analysis.")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.runs[0].italic = True
    caption.runs[0].font.color.rgb = RGBColor.from_string(GRAY)
    caption.runs[0].font.size = Pt(9)

    # Ensure Word refreshes page-number fields on open.
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")

    doc.core_properties.title = "PRISMA evidence-selection counts - prenatal chronic noise meta-analysis"
    doc.core_properties.subject = "Screening counts, full-text exclusion reasons, and PRISMA flowchart"
    doc.core_properties.author = "Anna Lenz"
    doc.save(path)


def main():
    _, included, excluded, not_retrieved, reason_counts = load_full_text_records()
    studies = load_analysis_studies()
    create_flowchart(PNG_OUT)
    create_docx(DOCX_OUT, PNG_OUT, included, excluded, not_retrieved, reason_counts, studies)
    print(f"Created {PNG_OUT}")
    print(f"Created {DOCX_OUT}")
    print(f"Reasons: {reason_counts}")


if __name__ == "__main__":
    main()
