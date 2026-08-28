from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path("/Users/annalenz/Desktop/ChronicNoise_Rodents/Prenatal_Chronic_Noise")
SOURCE = ROOT / "protocol/PCNE_protocol_deviation_offspring_age_class_2026-08-18.docx"
OUTPUT = ROOT / "protocol/PCNE_protocol_deviations_addendum_2026-08-28.docx"


def replace_paragraph_text_preserving_first_run(paragraph, text):
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run._element.getparent().remove(run._element)


def replace_label_value(paragraph, label, value):
    runs = paragraph.runs
    if len(runs) >= 2:
        runs[0].text = label
        runs[1].text = value
        for run in runs[2:]:
            run._element.getparent().remove(run._element)
        return
    replace_paragraph_text_preserving_first_run(paragraph, label + value)
    paragraph.runs[0].bold = True
    paragraph.add_run(value)


def add_labeled_paragraph_before(anchor, label, text):
    paragraph = anchor.insert_paragraph_before(style="Normal")
    run = paragraph.add_run(label)
    run.bold = True
    paragraph.add_run(text)
    return paragraph


doc = Document(SOURCE)

# Broaden the existing addendum so it can hold more than one approved deviation.
replace_paragraph_text_preserving_first_run(doc.paragraphs[0], "PROTOCOL DEVIATIONS ADDENDUM")
replace_paragraph_text_preserving_first_run(
    doc.paragraphs[1], "Post-extraction classification and coding decisions"
)
replace_label_value(doc.paragraphs[3], "Decision dates: ", "18 and 28 August 2026")
replace_label_value(
    doc.paragraphs[4],
    "Status: ",
    "Approved for the v10 codebook, data extraction, and analysis preparation",
)
replace_label_value(
    doc.paragraphs[5],
    "Scope: ",
    "Post-extraction classification and coding rules; no change to study eligibility or author-reported data",
)

section = doc.sections[0]
if section.header.paragraphs and section.header.paragraphs[0].runs:
    section.header.paragraphs[0].runs[0].text = "PCNE PROTOCOL DEVIATIONS ADDENDUM"
if section.footer.paragraphs and section.footer.paragraphs[0].runs:
    section.footer.paragraphs[0].runs[0].text = "PCNE | 2026-08-28 | Page "

# Insert the new decision immediately before the references section.
references_heading = next(
    p for p in doc.paragraphs if p.text.strip() == "5. References retained for citation"
)

heading = references_heading.insert_paragraph_before(
    "5. Control-condition uncertainty (Unclear)", style="Heading 1"
)

paragraph = references_heading.insert_paragraph_before(style="Normal")
paragraph.add_run("Registered definition: ").bold = True
paragraph.add_run(
    "The protocol defines control_conditions as the acoustic conditions under which control animals were kept during the maternal noise-exposure phase: silence, ambient sound, white noise, or other. It does not specify how to code a study when the control acoustic environment cannot be determined from the report."
)

paragraph = references_heading.insert_paragraph_before(style="Normal")
paragraph.add_run("Approved deviation: ").bold = True
paragraph.add_run(
    "Add Unclear as an allowed value for control_conditions. Code Unclear when the article identifies a control or no-noise group but provides insufficient information to distinguish silence from ambient sound or another acoustic condition. This includes descriptions such as “no noise exposure,” “without audiogenic stimulation,” or equivalent handling in an exposure room or chamber when its background sound or acoustic isolation is not reported."
)

paragraph = references_heading.insert_paragraph_before(style="Normal")
paragraph.add_run("Do not use Unclear when: ").bold = True
paragraph.add_run(
    "the source supports a more specific classification. Code ambient_sound when controls remained under stated standard or normal laboratory housing without experimental playback; code silence when controls were explicitly kept in an acoustically isolated, quiet, or noiseless environment without playback; code white_noise when white noise was explicitly used as the control sound; and code other for a reported acoustic condition that does not fit the preceding categories."
)

paragraph = references_heading.insert_paragraph_before(style="Normal")
paragraph.add_run("Documentation rule: ").bold = True
paragraph.add_run(
    "For every Unclear value, control_notes must record the authors’ description and the missing information that prevents a more specific classification. Leave control_loudness blank unless the control sound level was reported. Do not infer silence solely from the absence of experimental noise playback."
)

paragraph = references_heading.insert_paragraph_before(style="Normal")
paragraph.add_run("Rationale and impact: ").bold = True
paragraph.add_run(
    "This category prevents uncertain reports from being forced into silence or ambient sound and makes the uncertainty available for descriptive summaries and moderator analyses. It does not broaden the comparator eligibility criterion: eligible controls must still represent dams not exposed to the experimental chronic-noise treatment. The change harmonizes the registered definition with the operational moderator dictionary and the reviewed included-study evidence."
)

replace_paragraph_text_preserving_first_run(
    references_heading, "6. References retained for citation"
)

# Keep the new heading with its first paragraph and avoid orphaned labels.
for paragraph in doc.paragraphs:
    p_pr = paragraph._p.get_or_add_pPr()
    if paragraph.style.name.startswith("Heading"):
        keep_next = p_pr.find(qn("w:keepNext"))
        if keep_next is None:
            keep_next = OxmlElement("w:keepNext")
            p_pr.append(keep_next)

doc.core_properties.title = "PCNE protocol deviations addendum"
doc.core_properties.subject = (
    "Offspring age classification and control-condition uncertainty coding"
)
doc.save(OUTPUT)
print(OUTPUT)
