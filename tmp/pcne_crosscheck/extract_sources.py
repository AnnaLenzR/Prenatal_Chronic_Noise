from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
import pdfplumber
from docx import Document


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "tmp" / "pcne_crosscheck"


def clean(value):
    if pd.isna(value):
        return None
    if hasattr(value, "isoformat"):
        try:
            return value.isoformat()
        except Exception:
            pass
    if isinstance(value, (int, float, str, bool)):
        return value
    return str(value)


def extract_protocol() -> None:
    path = ROOT / "protocol" / "PCNE_protocol_v.4_111425.docx"
    doc = Document(path)
    lines: list[str] = []
    for i, paragraph in enumerate(doc.paragraphs, start=1):
        text = paragraph.text.strip()
        if text:
            lines.append(f"P{i:04d}\t{text}")
    for ti, table in enumerate(doc.tables, start=1):
        lines.append(f"\nTABLE {ti}")
        for ri, row in enumerate(table.rows, start=1):
            cells = [cell.text.replace("\n", " | ").strip() for cell in row.cells]
            lines.append(f"T{ti:02d}R{ri:03d}\t" + "\t".join(cells))
    (OUT / "protocol_extracted.txt").write_text("\n".join(lines), encoding="utf-8")


def extract_pdfs() -> list[dict]:
    manifest: list[dict] = []
    for path in sorted((ROOT / "included_studies").rglob("*.pdf")):
        rel = path.relative_to(ROOT).as_posix()
        pages: list[str] = []
        with pdfplumber.open(path) as pdf:
            for page_number, page in enumerate(pdf.pages, start=1):
                text = page.extract_text(x_tolerance=2, y_tolerance=3) or ""
                pages.append(f"\n===== PAGE {page_number} =====\n{text}")
            page_count = len(pdf.pages)
        out_name = rel.replace("/", "__").replace(".pdf", ".txt")
        (OUT / out_name).write_text("\n".join(pages), encoding="utf-8")
        manifest.append({"source_file": rel, "text_file": out_name, "pages": page_count})
    return manifest


def profile_excel_files() -> list[dict]:
    result: list[dict] = []
    for path in sorted((ROOT / "included_studies").rglob("*.xlsx")):
        rel = path.relative_to(ROOT).as_posix()
        xls = pd.ExcelFile(path)
        for sheet_name in xls.sheet_names:
            frame = pd.read_excel(path, sheet_name=sheet_name, header=None)
            preview = [[clean(v) for v in row] for row in frame.iloc[:20, :20].values.tolist()]
            result.append(
                {
                    "source_file": rel,
                    "sheet": sheet_name,
                    "rows": int(frame.shape[0]),
                    "columns": int(frame.shape[1]),
                    "preview": preview,
                }
            )
    return result


def profile_csv() -> dict:
    path = ROOT / "Data" / "PCN_data_ext_checking_v.3.csv"
    frame = pd.read_csv(path)
    profile = {
        "rows": int(frame.shape[0]),
        "columns": list(frame.columns),
        "study_counts": frame["study_id"].value_counts(dropna=False).to_dict(),
        "unique_counts": {col: int(frame[col].nunique(dropna=True)) for col in frame.columns},
        "missing_counts": {col: int(frame[col].isna().sum()) for col in frame.columns},
    }
    return profile


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    extract_protocol()
    manifest = {
        "pdfs": extract_pdfs(),
        "excel_sheets": profile_excel_files(),
        "csv": profile_csv(),
    }
    (OUT / "source_manifest.json").write_text(json.dumps(manifest, indent=2), encoding="utf-8")


if __name__ == "__main__":
    main()
